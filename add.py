import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
import os
import zipfile
import io
from datetime import datetime
import locale

# Configurar localização para português do Brasil
try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')
    except:
        pass  # Fallback para a localização padrão se pt_BR não estiver disponível

# Função para formatar o número do ofício com o ano atual
def formatar_numero_oficio(numero):
    ano_atual = datetime.now().year
    return f"{numero}/{ano_atual}"

# Função para obter data formatada em português do Brasil
def formatar_data_ptbr(data):
    try:
        # Tentativa de formatar com locale
        return data.strftime("%d de %B de %Y").lower()
    except:
        # Fallback para manual mapping se o locale não funcionar
        meses = {
            1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
            5: "maio", 6: "junho", 7: "julho", 8: "agosto",
            9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
        }
        return f"{data.day} de {meses[data.month]} de {data.year}"

# Função para formatar o documento
def formatar_documento(doc):
    # Estilo para todo o documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Configurar margens (em cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    return doc

# Função para adicionar parágrafo com espaçamento controlado
def adicionar_paragrafo(doc, texto, bold=False, italic=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, espacamento_antes=0, espacamento_depois=0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment
    run = paragraph.add_run(texto)
    run.bold = bold
    run.italic = italic
    
    # Ajustar espaçamento se fornecido
    if espacamento_antes > 0 or espacamento_depois > 0:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(espacamento_antes)
        paragraph_format.space_after = Pt(espacamento_depois)
        
    return paragraph

# Função para criar ofício modelo de comunicação de arquivamento
def criar_oficio_arquivamento(numero_oficio, data, numero_idea):
    doc = Document()
    doc = formatar_documento(doc)
    
    # Número do ofício com ano atual - ALINHADO À ESQUERDA
    adicionar_paragrafo(doc, f"OFÍCIO Nº {formatar_numero_oficio(numero_oficio)}/SP-FSA/25ªPJ", 
                    bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, espacamento_depois=6)
    
    # Referência IDEA - ALINHADO À ESQUERDA
    adicionar_paragrafo(doc, f"(Ref.: IDEA nº {numero_idea}/{datetime.now().year})", 
                    bold=True, italic=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, espacamento_depois=12)
    
    # Local e data - ALINHADO À DIREITA
    adicionar_paragrafo(doc, f"Feira de Santana, {data}", alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, espacamento_depois=12)
    
    # Destinatário (sem espaçamento entre linhas)
    adicionar_paragrafo(doc, "A Sua Excelência a Senhora", espacamento_depois=0)
    adicionar_paragrafo(doc, "MARIA CLÉCIA VASCONCELOS DE MORAIS FIRMINO COSTA", bold=True, espacamento_depois=0)
    adicionar_paragrafo(doc, "Delegacia Especializada de Atendimento à Mulher de Feira de Santana --", espacamento_depois=0)
    adicionar_paragrafo(doc, "DEAM", espacamento_depois=0)
    adicionar_paragrafo(doc, "Avenida Maria Quitéria nº 1870, Centro", espacamento_depois=0)
    adicionar_paragrafo(doc, "Feira de Santana -- Bahia, CEP: 44001-344", espacamento_depois=0)
    adicionar_paragrafo(doc, "E-mail: deam.feiradesantana@pcivil.ba.gov.br", espacamento_depois=12)
    
    # Vocativo
    adicionar_paragrafo(doc, "Excelentíssima Senhora,", espacamento_depois=12)
    
    # Conteúdo - JUSTIFICADO
    conteudo = (
        "Com os nossos cordiais cumprimentos, DE ORDEM DE DRA. NAYARA VALTÉRCIA "
        "GONÇALVES BARRETO, Promotora de Justiça titular da 25ª Promotoria de "
        "Justiça, sirvo-me do presente para, atendendo ao quanto disposto no art. "
        "28 do Código de Processo Penal, comunicar a Vossa Excelência o "
        f"ARQUIVAMENTO do Inquérito Policial IDEA nº {numero_idea}/{datetime.now().year}, consoante "
        "Promoção anexa."
    )
    adicionar_paragrafo(doc, conteudo, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, espacamento_depois=12)
    
    # Despedida
    adicionar_paragrafo(doc, "Cordialmente,", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, espacamento_depois=18)
    
    # Assinatura - CENTRALIZADA sem espaçamento
    adicionar_paragrafo(doc, "(assinado eletronicamente)", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, espacamento_depois=0)
    adicionar_paragrafo(doc, "ANDERSON MELO FIUSA BASTOS", bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, espacamento_depois=0)
    adicionar_paragrafo(doc, "Secretaria Processual", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    # Criar arquivo temporário
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    
    return temp_file.name

# Função para criar ofício de notificação para vítima (ofício 2)
def criar_oficio_notificacao_vitima(numero_oficio, data, numero_idea, nome_vitima, endereco, telefone):
    doc = Document()
    doc = formatar_documento(doc)
    
    # Número do ofício - ALINHADO À ESQUERDA
    adicionar_paragrafo(doc, f"OFÍCIO Nº {formatar_numero_oficio(numero_oficio)}/SP-FSA/25ªPJ", 
                    bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, espacamento_depois=6)

    # Referência IDEA - ALINHADO À ESQUERDA
    adicionar_paragrafo(doc, f"(Ref.: IDEA nº {numero_idea}/{datetime.now().year})", 
                    bold=True, italic=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, espacamento_depois=12)
    
    # Local e data - ALINHADO À DIREITA
    adicionar_paragrafo(doc, f"Feira de Santana, {data}", alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, espacamento_depois=12)
    
    # Destinatário (sem espaçamento entre linhas)
    adicionar_paragrafo(doc, "A Sua Senhoria", espacamento_depois=0)
    adicionar_paragrafo(doc, f"{nome_vitima}", espacamento_depois=0)
    adicionar_paragrafo(doc, f"{endereco}", espacamento_depois=0)
    
    if telefone:
        adicionar_paragrafo(doc, f"Tel: {telefone}", espacamento_depois=0)
    
    # Adicionar espaçamento após o bloco de destinatário
    adicionar_paragrafo(doc, "", espacamento_depois=12)
    
    # Vocativo
    adicionar_paragrafo(doc, "Ilustríssima Senhora,", espacamento_depois=12)
    
    # Conteúdo fixo para o ofício 2 - JUSTIFICADO
    conteudo = (
        "Com os nossos cordiais cumprimentos, DE ORDEM DE DRA. NAYARA VALTÉRCIA "
        "GONÇALVES BARRETO, Promotora de Justiça titular da 25ª Promotoria de "
        "Justiça de Feira de Santana, sirvo-me do presente para Notificá-la acerca "
        f"do ARQUIVAMENTO do Inquérito Policial IDEA nº {numero_idea}/{datetime.now().year}, "
        "no qual a Vossa Senhoria figura como vítima, consoante Promoção anexa."
    )
    adicionar_paragrafo(doc, conteudo, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, espacamento_depois=12)
    
    conteudo2 = (
        "Em não concordando com o arquivamento do expediente criminal em questão, "
        "poderá, no prazo de 30 (trinta) dias a contar do recebimento do presente, "
        "encaminhar recurso dirigido à Procuradoria-Geral de Justiça, nos termos "
        "do art. 28, §1º, do Código de Processo Penal). Para tanto, recomendamos "
        "que procure orientação jurídica adequada para o exercício desse direito."
    )
    adicionar_paragrafo(doc, conteudo2, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, espacamento_depois=12)
    
    conteudo3 = (
        "Por fim, requer que a resposta, se for o caso, seja enviada, preferencialmente, "
        "por meio eletrônico para o endereço de e-mail: sp.feiradesantana@mpba.mp.br."
    )
    adicionar_paragrafo(doc, conteudo3, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, espacamento_depois=12)
    
    # Despedida
    adicionar_paragrafo(doc, "Atenciosamente,", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, espacamento_depois=18)
    
    # Assinatura - CENTRALIZADA sem espaçamento
    adicionar_paragrafo(doc, "(assinado eletronicamente)", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, espacamento_depois=0)
    adicionar_paragrafo(doc, "ANDERSON MELO FIUSA BASTOS", bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, espacamento_depois=0)
    adicionar_paragrafo(doc, "Secretaria Processual", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    # Criar arquivo temporário
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    
    return temp_file.name

# Função para criar ofício de notificação para o acusado (ofício 3)
def criar_oficio_notificacao_acusado(numero_oficio, data, numero_idea, nome_acusado, endereco, telefone):
    doc = Document()
    doc = formatar_documento(doc)
    
    # Número do ofício - ALINHADO À ESQUERDA
    adicionar_paragrafo(doc, f"OFÍCIO Nº {formatar_numero_oficio(numero_oficio)}/SP-FSA/25ªPJ", 
                    bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, espacamento_depois=6)

    # Referência IDEA - ALINHADO À ESQUERDA
    adicionar_paragrafo(doc, f"(Ref.: IDEA nº {numero_idea}/{datetime.now().year})", 
                    bold=True, italic=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, espacamento_depois=12)
    
    # Local e data - ALINHADO À DIREITA
    adicionar_paragrafo(doc, f"Feira de Santana, {data}", alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, espacamento_depois=12)
    
    # Destinatário (sem espaçamento entre linhas)
    adicionar_paragrafo(doc, "A Sua Senhoria", espacamento_depois=0)
    adicionar_paragrafo(doc, f"{nome_acusado}", espacamento_depois=0)
    adicionar_paragrafo(doc, f"{endereco}", espacamento_depois=0)
    
    if telefone:
        adicionar_paragrafo(doc, f"Tel: {telefone}", espacamento_depois=0)
    
    # Adicionar espaçamento após o bloco de destinatário
    adicionar_paragrafo(doc, "", espacamento_depois=12)
    
    # Vocativo
    adicionar_paragrafo(doc, "Ilustríssimo Senhor,", espacamento_depois=12)
    
    # Conteúdo fixo para o ofício 3 - JUSTIFICADO
    conteudo = (
        "Com os nossos cordiais cumprimentos, DE ORDEM DE DRA. NAYARA VALTÉRCIA "
        "GONÇALVES BARRETO, Promotora de Justiça titular da 25ª Promotoria de "
        "Justiça de Feira de Santana, sirvo-me do presente para Notificá-lo acerca "
        f"do ARQUIVAMENTO do Inquérito Policial IDEA Nº {numero_idea}/{datetime.now().year}."
    )
    adicionar_paragrafo(doc, conteudo, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, espacamento_depois=12)
    
    # Despedida
    adicionar_paragrafo(doc, "Atenciosamente,", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, espacamento_depois=18)
    
    # Assinatura - CENTRALIZADA sem espaçamento
    adicionar_paragrafo(doc, "(assinado eletronicamente)", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, espacamento_depois=0)
    adicionar_paragrafo(doc, "ANDERSON MELO FIUSA BASTOS", bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, espacamento_depois=0)
    adicionar_paragrafo(doc, "Secretaria Processual", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    # Criar arquivo temporário
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    
    return temp_file.name

# Função para criar um arquivo ZIP com os ofícios
def criar_zip_oficios(arquivos):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for nome_arquivo, caminho_arquivo in arquivos.items():
            # Adicionar cada arquivo ao ZIP com nome adequado
            with open(caminho_arquivo, "rb") as f:
                zip_file.writestr(f"{nome_arquivo}.docx", f.read())
    
    zip_buffer.seek(0)
    return zip_buffer

# Inicializar o estado da sessão para armazenar os dados dos ofícios
if "dados_oficios" not in st.session_state:
    st.session_state.dados_oficios = {}

# Interface do Streamlit
st.set_page_config(page_title="Gerador de Ofícios - MPBA", layout="wide")

st.title("Gerador de Ofícios Automáticos - MPBA")
st.write("Preencha os dados abaixo para gerar os ofícios.")

# Criar abas para cada ofício
tab1, tab2, tab3 = st.tabs([
    "Ofício 1 - Comunicação à Delegacia", 
    "Ofício 2 - Notificação à Vítima", 
    "Ofício 3 - Notificação ao Acusado"
])

with tab1:
    st.subheader("Dados do Ofício 1 - Comunicação de Arquivamento à Delegacia")
    
    with st.form("dados_oficio_1"):
        numero_oficio = st.text_input("Número do Ofício", "4886")
        data_oficio = st.date_input("Data do Ofício")
        idea_numero = st.text_input("Número IDEA (número do processo)", "596.9.489799")
        
        submit_1 = st.form_submit_button("Salvar Dados do Ofício 1")
    
    if submit_1:
        # Calcular os números sequenciais para os outros ofícios
        try:
            num_oficio_1 = int(numero_oficio)
            num_oficio_2 = num_oficio_1 + 1
            num_oficio_3 = num_oficio_1 + 2
        except ValueError:
            num_oficio_1 = numero_oficio
            num_oficio_2 = numero_oficio
            num_oficio_3 = numero_oficio
            st.warning("O número do ofício deve ser um número inteiro para sequência automática.")
        
        # Formatar a data em português do Brasil
        data_formatada = formatar_data_ptbr(data_oficio)
        
        st.session_state.dados_oficios["oficio_1"] = {
            "numero_oficio": numero_oficio,
            "data_oficio": data_formatada,
            "idea_numero": idea_numero
        }
        
        # Pré-configurar os números sequenciais para os outros ofícios
        if "oficio_2" not in st.session_state.dados_oficios:
            st.session_state.dados_oficios["oficio_2"] = {"numero_oficio": str(num_oficio_2)}
        else:
            st.session_state.dados_oficios["oficio_2"]["numero_oficio"] = str(num_oficio_2)
            
        if "oficio_3" not in st.session_state.dados_oficios:
            st.session_state.dados_oficios["oficio_3"] = {"numero_oficio": str(num_oficio_3)}
        else:
            st.session_state.dados_oficios["oficio_3"]["numero_oficio"] = str(num_oficio_3)
            
        st.success("Dados do Ofício 1 salvos!")

with tab2:
    st.subheader("Dados do Ofício 2 - Notificação à Vítima")
    
    # Auto-preencher o número sequencial se estiver disponível
    numero_oficio_2_default = ""
    if "oficio_2" in st.session_state.dados_oficios and "numero_oficio" in st.session_state.dados_oficios["oficio_2"]:
        numero_oficio_2_default = st.session_state.dados_oficios["oficio_2"]["numero_oficio"]
    
    with st.form("dados_oficio_2"):
        numero_oficio_2 = st.text_input("Número do Ofício", numero_oficio_2_default)
        nome_vitima = st.text_input("Nome da Vítima")
        endereco_vitima = st.text_input("Endereço da Vítima")
        telefone_vitima = st.text_input("Telefone da Vítima (opcional)")
        
        # O conteúdo é fixo conforme solicitado
        st.info("O conteúdo deste ofício é padrão conforme a especificação.")
        
        submit_2 = st.form_submit_button("Salvar Dados do Ofício 2")
    
    if submit_2:
        st.session_state.dados_oficios["oficio_2"] = {
            "numero_oficio": numero_oficio_2,
            "nome": nome_vitima,
            "endereco": endereco_vitima,
            "telefone": telefone_vitima
        }
        st.success("Dados do Ofício 2 salvos!")

with tab3:
    st.subheader("Dados do Ofício 3 - Notificação ao Acusado")
    
    # Auto-preencher o número sequencial se estiver disponível
    numero_oficio_3_default = ""
    if "oficio_3" in st.session_state.dados_oficios and "numero_oficio" in st.session_state.dados_oficios["oficio_3"]:
        numero_oficio_3_default = st.session_state.dados_oficios["oficio_3"]["numero_oficio"]
    
    with st.form("dados_oficio_3"):
        numero_oficio_3 = st.text_input("Número do Ofício", numero_oficio_3_default)
        nome_acusado = st.text_input("Nome do Acusado")
        endereco_acusado = st.text_input("Endereço do Acusado")
        telefone_acusado = st.text_input("Telefone do Acusado (opcional)")
        
        # O conteúdo é fixo conforme solicitado
        st.info("O conteúdo deste ofício é padrão conforme a especificação.")
        
        submit_3 = st.form_submit_button("Salvar Dados do Ofício 3")
    
    if submit_3:
        st.session_state.dados_oficios["oficio_3"] = {
            "numero_oficio": numero_oficio_3,
            "nome": nome_acusado,
            "endereco": endereco_acusado,
            "telefone": telefone_acusado
        }
        st.success("Dados do Ofício 3 salvos!")

# Mostrar o status atual dos dados salvos
status_col1, status_col2, status_col3 = st.columns(3)
with status_col1:
    if "oficio_1" in st.session_state.dados_oficios:
        st.info(f"Ofício 1: Dados do Ofício nº {formatar_numero_oficio(st.session_state.dados_oficios['oficio_1']['numero_oficio'])} salvos ✅")
    else:
        st.warning("Ofício 1: Dados não salvos ❌")
        
with status_col2:
    if "oficio_2" in st.session_state.dados_oficios:
        nome_info = st.session_state.dados_oficios["oficio_2"].get("nome", "")
        numero_info = st.session_state.dados_oficios["oficio_2"].get("numero_oficio", "")
        st.info(f"Ofício 2: Nº {formatar_numero_oficio(numero_info)} - {nome_info} ✅")
    else:
        st.warning("Ofício 2: Dados não salvos ❌")
        
with status_col3:
    if "oficio_3" in st.session_state.dados_oficios:
        nome_info = st.session_state.dados_oficios["oficio_3"].get("nome", "")
        numero_info = st.session_state.dados_oficios["oficio_3"].get("numero_oficio", "")
        st.info(f"Ofício 3: Nº {formatar_numero_oficio(numero_info)} - {nome_info} ✅")
    else:
        st.warning("Ofício 3: Dados não salvos ❌")

# Botão para gerar todos os ofícios
if st.button("Gerar Todos os Ofícios"):
    if not all(f"oficio_{i}" in st.session_state.dados_oficios for i in range(1, 4)):
        st.warning("Preencha e salve os dados de todos os três ofícios antes de gerar!")
        # Mostrar quais ofícios estão faltando
        for i in range(1, 4):
            if f"oficio_{i}" not in st.session_state.dados_oficios:
                st.warning(f"Ofício {i} não tem dados salvos!")
    else:
        st.success("Gerando ofícios...")
        
        # Obter o número IDEA do ofício 1 (será usado em todos os ofícios)
        idea_numero = st.session_state.dados_oficios["oficio_1"]["idea_numero"]
        
        # Criar os três documentos
        arquivos = {
            f"Ofício {st.session_state.dados_oficios['oficio_1']['numero_oficio']} - Comunicação à Delegacia": criar_oficio_arquivamento(
                st.session_state.dados_oficios["oficio_1"]["numero_oficio"],
                st.session_state.dados_oficios["oficio_1"]["data_oficio"],
                idea_numero
            ),
            f"Ofício {st.session_state.dados_oficios['oficio_2']['numero_oficio']} - Notificação à Vítima": criar_oficio_notificacao_vitima(
                st.session_state.dados_oficios["oficio_2"]["numero_oficio"],
                st.session_state.dados_oficios["oficio_1"]["data_oficio"],  # Usamos a mesma data do ofício 1
                idea_numero,
                st.session_state.dados_oficios["oficio_2"]["nome"],
                st.session_state.dados_oficios["oficio_2"]["endereco"],
                st.session_state.dados_oficios["oficio_2"].get("telefone", "")
            ),
            f"Ofício {st.session_state.dados_oficios['oficio_3']['numero_oficio']} - Notificação ao Acusado": criar_oficio_notificacao_acusado(
                st.session_state.dados_oficios["oficio_3"]["numero_oficio"],
                st.session_state.dados_oficios["oficio_1"]["data_oficio"],  # Usamos a mesma data do ofício 1
                idea_numero,
                st.session_state.dados_oficios["oficio_3"]["nome"],
                st.session_state.dados_oficios["oficio_3"]["endereco"],
                st.session_state.dados_oficios["oficio_3"].get("telefone", "")
            )
        }
        
        # Criar um ZIP com todos os arquivos
        zip_buffer = criar_zip_oficios(arquivos)
        
        # Botão para baixar o ZIP
        st.download_button(
            label="Baixar Todos os Ofícios (ZIP)",
            data=zip_buffer,
            file_name="Oficios.zip",
            mime="application/zip",
            key="download-zip"
        )
        
        # Opção para baixar individualmente também
        st.subheader("Ou baixe individualmente:")
        download_col1, download_col2, download_col3 = st.columns(3)
        
        keys = list(arquivos.keys())
        
        with download_col1:
            with open(arquivos[keys[0]], "rb") as file:
                st.download_button(
                    label=f"Baixar {keys[0]}",
                    data=file,
                    file_name=f"{keys[0]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download-1"
                )
        
        with download_col2:
            with open(arquivos[keys[1]], "rb") as file:
                st.download_button(
                    label=f"Baixar {keys[1]}",
                    data=file,
                    file_name=f"{keys[1]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download-2"
                )
        
        with download_col3:
            with open(arquivos[keys[2]], "rb") as file:
                st.download_button(
                    label=f"Baixar {keys[2]}",
                    data=file,
                    file_name=f"{keys[2]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download-3"
                )
        
        # Remover arquivos temporários após um tempo razoável para download
        for caminho in arquivos.values():
            try:
                os.remove(caminho)
            except:
                pass

# Adicionar um botão para limpar todos os dados
if st.button("Limpar Todos os Dados"):
    st.session_state.dados_oficios = {}
    st.success("Todos os dados foram limpos!")
    st.experimental_rerun()
